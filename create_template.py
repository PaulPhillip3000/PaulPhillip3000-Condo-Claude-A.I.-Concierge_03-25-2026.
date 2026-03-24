from docx import Document

def create_base_template():
    doc = Document()
    
    doc.add_paragraph("{{ ai_generated_statutory_text }}")
    
    doc.add_paragraph("Itemized Ledger of Amounts Owed:")
    
    # DocxTemplate loops syntax in tables:
    # {% tr for item in ledger %}
    # {{ item.Date }} | {{ item.Description }} | {{ item.Amount }}
    # {% tr endfor %}
    
from docx import Document

def create_base_template():
    doc = Document()
    
    doc.add_paragraph("{{ ai_generated_statutory_text }}")
    
    doc.add_paragraph("Itemized Ledger of Amounts Owed:")
    
    # Simpler text-based layout for the ledger to avoid xml/table tag issues with docxtpl
    doc.add_paragraph("{% for item in ledger %}{{ item.Date }} \t {{ item.Description }} \t ${{ item.Amount }}{% endfor %}")
    
    doc.add_paragraph("\nTotal Due: ${{ total_due }}")
    
    doc.save("template.docx")
    print("Created template.docx successfully.")

if __name__ == "__main__":
    create_base_template()
    print("Created template.docx successfully.")

if __name__ == "__main__":
    create_base_template()
